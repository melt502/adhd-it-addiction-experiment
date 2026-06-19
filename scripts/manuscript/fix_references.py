from pathlib import Path
from docx import Document

src = Path(r'D:/essay/adhd-1/LLM_ADHD612_AMPPS_stats_PDU_Lin_revised(2)_data_analysis_completed_final_author_polished_ampps_revised.docx')
out = Path(r'D:/essay/adhd-1/LLM_ADHD612_AMPPS_stats_PDU_Lin_revised(2)_data_analysis_completed_final_author_polished_ampps_references_fixed.docx')
doc = Document(src)

replacements = {
155: 'Aher, G. V., Arriaga, R. I., & Kalai, A. T. (2023). Using large language models to simulate multiple humans and replicate human subject studies. In Proceedings of the 40th International Conference on Machine Learning. [Proceedings volume, page range, and persistent URL require author verification.]',
156: 'Argyle, L. P., Busby, E. C., Fulda, N., Gubler, J. R., Rytting, C., & Wingate, D. (2023). Out of one, many: Using language models to simulate human samples. Political Analysis, 31(3), 337-351. https://doi.org/10.1017/pan.2023.2',
157: 'Augner, C., Vlasak, T., & Barth, A. (2023). The relationship between problematic internet use and attention deficit, hyperactivity and impulsivity: A meta-analysis. Journal of Psychiatric Research, 168, 1-12. https://doi.org/10.1016/j.jpsychires.2023.10.032',
158: 'Bai, X., Wang, A., Sucholutsky, I., & Griffiths, T. L. (2025). Explicitly unbiased large language models still form biased associations. Proceedings of the National Academy of Sciences, 122(8), Article e2416228122. https://doi.org/10.1073/pnas.2416228122',
159: 'Binz, M., Akata, E., Bethge, M., Brändle, F., Callaway, F., Coda-Forno, J., et al. (2025). A foundation model to predict and capture human cognition. Nature, 644(8078), 1002-1009. https://doi.org/10.1038/s41586-025-09215-4',
160: 'Callan, P. D., Swanberg, S., Weber, S. K., Eidnes, K., Pope, T. M., & Shepler, D. (2024). Diagnostic utility of Conners Continuous Performance Test-3 for attention deficit/hyperactivity disorder: A systematic review. Journal of Attention Disorders, 28(6), 992-1007. https://doi.org/10.1177/10870547231223727',
161: 'Chen, S. H., Weng, L. J., Su, Y. J., Wu, H. M., & Yang, P. F. (2003). Development of a Chinese Internet addiction scale and its psychometric study. Chinese Journal of Psychology, 45(3), 279-294. [DOI or persistent URL requires author verification.]',
162: 'Chiappone, A., et al. (2026). LLM-based neurodivergent psychometric profile simulation. Preprint. [Full author list, title, repository, version, and DOI or persistent URL require author verification.]',
163: "Conners, C. K., Erhardt, D., & Sparrow, E. (2002). Conners' Adult ADHD Rating Scales, Self-Report: Short Version (CAARS-S:S). Multi-Health Systems.",
164: 'Cui, Z., Li, N., & Zhou, H. (2025). A large-scale replication of scenario-based experiments in psychology and management using large language models. Nature Computational Science, 5(8), 627-634. https://doi.org/10.1038/s43588-025-00840-7',
165: 'Dillion, D., Tandon, N., Gu, Y., & Gray, K. (2023). Can AI language models replace human participants? Trends in Cognitive Sciences, 27(7), 597-600. https://doi.org/10.1016/j.tics.2023.04.008',
166: 'Gao, Y., Lee, H. R., Burtch, G., & Fazelpour, S. (2025). Take caution in using LLMs as human surrogates. Proceedings of the National Academy of Sciences, 122(24), Article e2501660122. https://doi.org/10.1073/pnas.2501660122',
167: 'Gonnermann-Müller, J., Haase, J., Leins, N., Kosch, T., & Pokutta, S. (2026). Maintaining stable personas? Examining temporal stability in LLM-based human simulation. In Proceedings of the Extended Abstracts of the 2026 CHI Conference on Human Factors in Computing Systems (pp. 1-6). https://doi.org/10.1145/3772363.3799334',
168: 'Harrison, A. G., Nay, S., & Armstrong, I. T. (2019). Diagnostic accuracy of the Conners’ Adult ADHD Rating Scale in a postsecondary population. Journal of Attention Disorders, 23(14), 1829-1837. https://doi.org/10.1177/1087054715625299',
169: 'Horton, J. J. (2023). Large language models as simulated economic agents: What can we learn from Homo silicus? SSRN Electronic Journal. https://doi.org/10.2139/ssrn.4413859',
170: 'Hu, T., Kyrychenko, Y., Rathje, S., Collier, N., van der Linden, S., & Roozenbeek, J. (2025). Generative language models exhibit social identity biases. Nature Computational Science, 5(1), 65-75. https://doi.org/10.1038/s43588-024-00741-1',
171: 'Humphreys, K. L., & Lee, S. S. (2011). Risk taking and sensitivity to punishment in children with ADHD, ODD, ADHD+ODD, and controls. Journal of Psychopathology and Behavioral Assessment, 33(3), 299-307. https://doi.org/10.1007/s10862-011-9237-6',
172: 'Jackson, J. N. S., & MacKillop, J. (2016). Attention-deficit/hyperactivity disorder and monetary delay discounting: A meta-analysis of case-control studies. Biological Psychiatry: Cognitive Neuroscience and Neuroimaging, 1(4), 316-325. https://doi.org/10.1016/j.bpsc.2016.01.007',
173: 'Kofler, M. J., Irwin, L. N., Soto, E. F., Groves, N. B., Harmon, S. L., & Sarver, D. E. (2019). Executive functioning heterogeneity in pediatric ADHD. Journal of Abnormal Child Psychology, 47(2), 273-286. https://doi.org/10.1007/s10802-018-0438-2',
174: 'Kuss, D. J., & Lopez-Fernandez, O. (2016). Internet addiction and problematic Internet use: A systematic review of clinical research. World Journal of Psychiatry, 6(1), 143-176. https://doi.org/10.5498/wjp.v6.i1.143',
175: 'Lejuez, C. W., Read, J. P., Kahler, C. W., Richards, J. B., Ramsey, S. E., Stuart, G. L., Strong, D. R., & Brown, R. A. (2002). Evaluation of a behavioral measure of risk taking: The Balloon Analogue Risk Task (BART). Journal of Experimental Psychology: Applied, 8(2), 75-84. https://doi.org/10.1037/1076-898X.8.2.75',
176: 'Lin, Z. (2025). Six fallacies in substituting large language models for human participants. PsyArXiv. https://doi.org/10.31234/osf.io/uqxcb_v2',
177: 'Lin, Z. (2026). Large language models as psychological simulators: A methodological guide. Advances in Methods and Practices in Psychological Science, 9(1). https://doi.org/10.1177/25152459251410153',
178: 'Salecha, A., Ireland, M. E., Subrahmanya, S., Sedoc, J., Ungar, L. H., & Eichstaedt, J. C. (2024). Large language models display human-like social desirability biases in Big Five personality surveys. PNAS Nexus, 3(12), Article pgae533. https://doi.org/10.1093/pnasnexus/pgae533',
179: 'Schröder, S., Morgenroth, T., Kuhl, U., Vaquet, V., & Paaßen, B. (2025). Large language models do not simulate human psychology. Preprint. [Repository, version, and DOI or persistent URL require author verification.]',
180: 'Sonuga-Barke, E. J. S. (2003). The dual pathway model of AD/HD: An elaboration of neuro-developmental characteristics. Neuroscience & Biobehavioral Reviews, 27(7), 593-604. https://doi.org/10.1016/j.neubiorev.2003.08.005',
181: 'Varela, J. L., Magnante, A. T., Miskey, H. M., Ord, A. S., Eldridge, A., & Shura, R. D. (2024). A systematic review of the utility of continuous performance tests among adults with ADHD. The Clinical Neuropsychologist, 38(7), 1524-1585. https://doi.org/10.1080/13854046.2024.2315740',
182: 'Wang, A., Morgenstern, J., & Dickerson, J. P. (2025). Large language models that replace human participants can harmfully misportray and flatten identity groups. Nature Machine Intelligence, 7(3), 400-411. https://doi.org/10.1038/s42256-025-00986-z',
183: 'Willcutt, E. G., Doyle, A. E., Nigg, J. T., Faraone, S. V., & Pennington, B. F. (2005). Validity of the executive function theory of attention-deficit/hyperactivity disorder: A meta-analytic review. Biological Psychiatry, 57(11), 1336-1346. https://doi.org/10.1016/j.biopsych.2005.02.006',
184: 'Wu, X. P., Sun, J. H., Li, Q. Q., & Guo, L. T. (2009). Reliability and validity of the Chinese version of Conners’ Adult ADHD Rating Scales Self-Report. Chinese Mental Health Journal, 23(5), 349-352. [DOI or persistent URL requires author verification.]',
185: 'Young, K. S. (1998). Internet addiction: The emergence of a new clinical disorder. CyberPsychology & Behavior, 1(3), 237-244. https://doi.org/10.1089/cpb.1998.1.237',
}

for i, text in replacements.items():
    p = doc.paragraphs[i]
    if p.runs:
        p.runs[0].text = text
        for run in p.runs[1:]:
            run.text = ''
    else:
        p.text = text

for p in doc.paragraphs:
    if 'Dillon et al.' in p.text:
        for run in p.runs:
            run.text = run.text.replace('Dillon et al.', 'Dillion et al.')
    if 'ADHD student-persona simulations show stable self-reports but weaker or drifting observable behavior' in p.text:
        for run in p.runs:
            run.text = run.text.replace('ADHD student-persona simulations show stable self-reports but weaker or drifting observable behavior', 'LLM-based human-simulation studies examine persona stability but raise questions about observable behavior')

doc.save(out)
print(out)
