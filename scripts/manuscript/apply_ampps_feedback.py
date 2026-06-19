from pathlib import Path
from docx import Document

src = Path(r'D:/essay/adhd-1/LLM_ADHD612_AMPPS_stats_PDU_Lin_revised(2)_data_analysis_completed_final_author_polished.docx')
out = Path(r'D:/essay/adhd-1/LLM_ADHD612_AMPPS_stats_PDU_Lin_revised(2)_data_analysis_completed_final_author_polished_ampps_revised.docx')

doc = Document(src)

replacements = {
    13: "Large language models (LLMs) are increasingly used as simulated participants, but their psychological plausibility may reflect stereotypes rather than valid behavioral prediction. We evaluated this risk using ADHD-related traits in young adults (questionnaire N = 2,016; task subsample N = 157). Human data showed robust associations between ADHD-related traits and problematic digital-use measures. Task behavior, however, was heterogeneous: CPT-defined I/C participants showed near-zero Stroop differences, modestly lower n-back accuracy, and lower, not higher, BART adjusted pumps. By contrast, five LLM families generated coherent ADHD-like profiles characterized by poorer executive-function performance, higher risk-taking, and stronger immediate-reward preference. Cue-contribution analyses showed that labels and symptom descriptions, rather than anonymized numeric scores, drove these outputs. Individual-level LLM predictions did not outperform statistical baselines. A secondary problematic digital-use generalization check showed similar cue dominance and effect-size inflation in reward-risk predictions. These findings indicate that LLM-simulated clinical or behavioral profiles require human benchmarks and cue-sensitivity audits before they are used as substitutes for participants.",
    28: "General disclosures. Conflicts of interest: The authors declare no conflicts of interest with respect to the authorship or publication of this article. Funding: This work was supported by the National Social Science Fund of China (22BSH093). Artificial intelligence: The authors used AI-assisted tools for programming assistance, formatting, language refinement, and manuscript organization. All data analyses, outputs, interpretations, and final manuscript decisions were checked and verified by the authors. The LLMs reported as study objects were not used as authors or autonomous analysts. Ethics: The study was reviewed and approved by the Ethics Committee of Central China Normal University. All participants provided informed consent before participation, and all analyzed data were de-identified. The final approval number and approval date should be verified by the authors before submission.",
    29: "Preregistration. The analyses were not preregistered. The manuscript distinguishes primary analyses from the secondary problematic digital-use extension. Deviations and exploratory analyses are documented in the Supplemental Material.",
    31: "Statistical reproducibility. Inferential statistics, effect-size estimates, confidence intervals, exclusion counts, and baseline-prediction metrics can be regenerated from the analysis scripts deposited in the repository. The manuscript reports exact p values where feasible, effect sizes in raw and standardized units, and 95% confidence intervals for effect sizes.",
    33: "Reporting. We report how the sample size was determined, data exclusions, prompt manipulations, and measures included in the study. Detailed run logs and exclusion logs are provided in the Supplemental Material.",
    41: "We used a convenience sample recruited in classroom settings at two universities in Hubei Province, China. The proximal population was young adults enrolled at these universities, and generalization should therefore be limited to similar young-adult educational samples unless replicated in broader populations. The survey yielded responses from 2,100 participants. We excluded responses with an average response time below 2 s per item, patterned response styles, or incomplete questionnaires. Screening retained 2,016 complete and valid responses, corresponding to a valid-completion rate of 96.0%. Demographic information included age, gender, residence, family structure, only-child status, and gaming experience. Participants had a mean age of 19.35 years; 733 identified as male and 1,283 as female. The available eligible sample determined the sample size. No prospective power analysis was preregistered for the secondary LLM-validation analyses. All participants provided informed consent before participation.",
    46: "The experimental subsample (N = 157) completed computerized Stroop, n-back, BART, DDT, and CPT-derived attention-profile assessments. Questionnaire and task records were linked using de-identified participant identifiers. CPT-derived labels were treated as attention-profile indicators rather than clinical diagnoses. Main analyses contrasted Normal versus I/C profiles, with uncertain labels handled cautiously or in sensitivity analyses. Accordingly, we refer to ADHD-related traits, CPT-defined attention profiles, and ADHD-like simulated profiles rather than clinically diagnosed ADHD patients.",
    47: "",
    53: "Note. The table summarizes model roles and run settings reported in the manuscript. Model versions, access dates, API or platform names, maximum-token settings, invalid-output counts, and component-specific output counts should be checked against the available repository logs and run manifests before submission. Full model-specific run counts are reported in the Supplemental Material.",
    61: "Cue-contribution analyses used standardized domain-level contrasts, aligned so that larger values indicated stronger stereotype-oriented output. Label, numeric-score, and symptom-language effects were estimated as average marginal cue contributions across factorial prompt conditions. For individual prediction, repeated LLM outputs for the same participant, task, model, and input condition were averaged before evaluation. Prediction validity was assessed using Pearson correlation, Spearman correlation, RMSE, and out-of-sample R² relative to the mean baseline. Statistical baselines included mean prediction, CPT/PDU group mean, ridge regression, random forest, and label-augmented variants. Baseline models were evaluated using participant-level cross-validation. All outcomes from the same participant were kept in the same fold to prevent leakage. Preprocessing and model tuning were performed within training folds. The exact fold structure, repeat count, and hyperparameter grids are reported in the analysis scripts and Supplemental Material.",
    92: "The factorial cue results were clear (Figure 3). Score-only prompts produced a weak ADHD-like contrast (mean z = 0.349), similar to the weak no-cue baseline (mean z = 0.491). By contrast, label-only (mean z = 2.005) and symptom-only (mean z = 2.061) prompts produced large stereotype-oriented shifts. Combined cue conditions also remained high: label+score (2.567), label+symptom (2.569), score+symptom (2.400), and full-profile (2.681). These values are standardized contrasts; values around 2 indicate that the model separated ADHD-like profiles from the low-ADHD baseline by roughly two standard-deviation units in the stereotype-oriented direction.",
    108: "Note. R² values are out-of-sample values relative to the sample-mean baseline. Negative values indicate performance worse than the mean baseline. Full model-specific prediction metrics are reported in the Supplemental Material.",
    121: "The ADHD literature is central to interpreting this dissociation. Executive-function, delay-aversion, and reward-processing theories support expected associations between ADHD-related traits and tasks such as Stroop, n-back, BART, and delay discounting. The same literature also emphasizes heterogeneity, pathway diversity, and measurement limits. LLM outputs often followed an overgeneralized textbook narrative of poorer executive control and greater impulsive risk-taking. By contrast, the empirical benchmark showed task-specific patterns, including the BART direction reversal in the ADHD benchmark and the effect-size inflation in the PDU extension. The BART result is therefore not merely a numerical anomaly. It shows how a semantically plausible clinical narrative can fail as a behavioral prediction.",
    126: "The problematic digital-use extension broadens the interpretation without changing the evidentiary hierarchy. It suggests that cue-driven stereotype amplification may not be specific to ADHD-like labels. Models exaggerated the association between high problematic digital-use cues and BART risk-taking despite negligible human group differences, and they did not provide strong individual-level prediction. Because this extension was not as uniformly complete across all models and conditions as the ADHD analyses, it should be interpreted as supportive evidence for generality rather than as an equally weighted primary study. Because the PDU extension was not the primary stress test and had more uneven condition coverage, we do not use it to make independent claims about problematic digital-use mechanisms.",
    133: "Conceptualization: Y.L. and W.Z. Methodology: Y.L., S.T., and W.Z. Data curation: W.Z. and C.Z. Formal analysis: S.T. and Y.L. Investigation: W.Z. and S.T. Writing, original draft: W.Z., S.T., C.Z., and Y.L. Writing, review and editing: W.Z., S.T., C.Z., and Y.L. Supervision: Y.L. and W.Z. Project administration: W.Z. Funding acquisition: W.Z. All authors must read and approve the final manuscript and agree to be accountable for the work.",
    139: "The authors used AI-assisted tools for programming assistance, formatting, language refinement, and manuscript organization. All data analyses, outputs, interpretations, and final manuscript decisions were checked and verified by the authors. The authors take full responsibility for the manuscript content. The LLMs reported as study objects were not used as authors or autonomous analysts.",
    141: "The Supplemental Material includes human data cleaning and task metrics, prompt templates and JSON schemas, model API settings and run logs, cue-contribution calculations, factorial-regression sensitivity analyses, full model outputs by task and profile, prompt-ablation results, individual-prediction metrics, sensitivity analyses, problematic digital-use extension tables, run-quality checks, supplementary tables, and the codebook and variable dictionary.",
    153: "No prior version of this manuscript has been posted or submitted elsewhere.",
}

for i, text in replacements.items():
    p = doc.paragraphs[i]
    if p.runs:
        p.runs[0].text = text
        for run in p.runs[1:]:
            run.text = ''
    else:
        p.text = text

for i in range(142, 152):
    p = doc.paragraphs[i]
    if p.runs:
        for run in p.runs:
            run.text = ''
    else:
        p.text = ''

ref_replacements = {
    'Binz, M., et al. (2025). A foundation model to predict and capture human cognition. Nature.': 'Binz, M., et al. (2025). A foundation model to predict and capture human cognition. Nature. [Full author list, volume, article/page number, and DOI require author verification.]',
    'Cui, Z., Li, N., & Zhou, H. (2024/2025). A large-scale replication of scenario-based experiments in psychology and management using large language models. Nature Computational Science.': 'Cui, Z., Li, N., & Zhou, H. (year to be verified). A large-scale replication of scenario-based experiments in psychology and management using large language models. Nature Computational Science. [Publication year, volume, article/page number, and DOI require author verification.]',
    'Dillion, D., Tandon, N., Gu, Y., & Gray, K. (2023). Can AI language models replace human participants? Trends in Cognitive Sciences.': 'Dillon, D., Tandon, N., Gu, Y., & Gray, K. (2023). Can AI language models replace human participants? Trends in Cognitive Sciences. [Volume, issue, page range, and DOI require author verification.]',
    'Gao, Y., et al. (2025). Take caution in using LLMs as human surrogates. Proceedings of the National Academy of Sciences / related commentary.': 'Gao, Y., et al. (2025). Take caution in using LLMs as human surrogates. [Source, full author list, volume, article/page number, and DOI require author verification.]',
    'Jackson, J. N. S., MacKillop, J., & others. (2016). Attention-deficit/hyperactivity disorder and monetary delay discounting: A meta-analysis of case-control studies. Clinical Psychology Review.': 'Jackson, J. N. S., MacKillop, J., et al. (2016). Attention-deficit/hyperactivity disorder and monetary delay discounting: A meta-analysis of case-control studies. Clinical Psychology Review. [Complete author list, volume, page range, and DOI require author verification.]',
    'Schröder, S., Morgenroth, T., Kuhl, U., Vaquet, V., & Paaßen, B. (2025). Large language models do not simulate human psychology. Preprint.': 'Schröder, S., Morgenroth, T., Kuhl, U., Vaquet, V., & Paaßen, B. (2025). Large language models do not simulate human psychology. Preprint. [Repository, version, and DOI or persistent URL require author verification.]',
    'Gonnermann-Müller, J., Haase, J., Leins, N., Kosch, T., & Pokutta, S. (2026). LLM-based educational simulation: Evaluating temporal student persona stability across ADHD profiles. Preprint.': 'Gonnermann-Müller, J., Haase, J., Leins, N., Kosch, T., & Pokutta, S. (2026). LLM-based educational simulation: Evaluating temporal student persona stability across ADHD profiles. Preprint. [Repository, version, and DOI or persistent URL require author verification.]',
    'Wu, X, P., Sun, J, H., Li, Q, Q., & Guo, L, T. (2009). Reliability and Validity of Chinese Version of Conners’ Adult ADHD Ratting Scales Self-Report. Chinese Mental Health Journal, 23(5), 349-352.': 'Wu, X. P., Sun, J. H., Li, Q. Q., & Guo, L. T. (2009). Reliability and validity of the Chinese version of Conners’ Adult ADHD Rating Scales Self-Report. Chinese Mental Health Journal, 23(5), 349-352.',
    'Young KS. (1998). Internet addiction: the emergence of a new clinical disorder. CyberPsychology and Behavior, 1(1).': 'Young, K. S. (1998). Internet addiction: The emergence of a new clinical disorder. CyberPsychology & Behavior, 1(3), 237-244. [DOI requires author verification.]',
}

for p in doc.paragraphs:
    text = p.text.strip()
    if text in ref_replacements:
        new = ref_replacements[text]
        if p.runs:
            p.runs[0].text = new
            for run in p.runs[1:]:
                run.text = ''
        else:
            p.text = new
    if 'Dillion et al.' in p.text:
        for run in p.runs:
            run.text = run.text.replace('Dillion et al.', 'Dillon et al.')

if doc.tables:
    table = doc.tables[0]
    for row in table.rows[1:4]:
        if row.cells[3].text.strip() == 'multiple experiments':
            row.cells[3].text = 'component-specific counts in Supplemental Material'

doc.save(out)
print(out)
